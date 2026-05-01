from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt


def generate_ieee_report(
    filename: str,
    input_excerpt: str,
    plagiarism_percentage: float,
    humanized_percentage: float,
    ai_generated_percentage: float,
    top_sources: list[dict],
) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    doc.add_heading("GH Plagiarism Checker - Analysis Report", level=0)
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(f"Automated analysis summary for {filename}.")
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Plagiarism uses semantic embeddings + cosine similarity. Human/AI score uses RoBERTa detectors."
    )
    doc.add_heading("Results", level=1)
    doc.add_paragraph(f"Plagiarism Percentage: {plagiarism_percentage:.2f}%")
    doc.add_paragraph(f"Humanized Percentage: {humanized_percentage:.2f}%")
    doc.add_paragraph(f"AI-Generated Percentage: {ai_generated_percentage:.2f}%")
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(f"Input excerpt: {input_excerpt[:400]}...")
    for i, source in enumerate(top_sources[:3], start=1):
        doc.add_paragraph(
            f"{i}) {source.get('title', source.get('filename'))} - {source.get('similarity', 0):.2f}%"
        )
        if source.get("url"):
            doc.add_paragraph(f"URL: {source['url']}")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream
