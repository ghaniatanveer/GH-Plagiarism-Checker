from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def set_ieee_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)


def add_section(doc: Document, title: str, body: str) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)


def build_report(output_path: Path) -> None:
    doc = Document()
    set_ieee_style(doc)

    doc.add_heading("GH Plagiarism Checker - IEEE Project Report", level=0)
    doc.add_paragraph(
        "This report documents the design and implementation of GH Plagiarism Checker, "
        "a full-stack web platform for text extraction, plagiarism analysis, AI-writing "
        "detection, and IEEE-style report generation."
    )

    add_section(
        doc,
        "Abstract",
        "GH Plagiarism Checker accepts user input as direct text or uploaded PDF/DOCX/TXT files, "
        "extracts clean text, computes a plagiarism percentage using semantic similarity, computes "
        "humanized and AI-detected percentages using transformer-based classification models, and "
        "returns downloadable analysis reports. The system is built with a React frontend and a "
        "FastAPI backend and is optimized for fast response in practical use.",
    )

    add_section(
        doc,
        "1. Introduction",
        "Academic integrity verification requires more than lexical matching. This project addresses "
        "that requirement by combining semantic embedding comparison with modern AI-text detection. "
        "The platform is designed for students, researchers, and reviewers who need quick topic-aware "
        "source matching and understandable percentages for plagiarism and AI assistance.",
    )

    add_section(
        doc,
        "2. System Architecture",
        "The frontend is implemented in React (Vite) and handles file upload, text input, progress "
        "visualization, and report download. The backend is implemented in FastAPI and exposes REST "
        "endpoints for analysis and report retrieval. Core backend modules include text extraction, "
        "plagiarism detection, human/AI detection, and Word report generation.",
    )

    add_section(
        doc,
        "3. Data Flow and Workflow",
        "Step 1: User submits text or document. Step 2: Backend extracts and normalizes text. "
        "Step 3: Plagiarism detector computes semantic similarity against open-source candidate "
        "documents and local fallback corpus. Step 4: Human detector computes AI-generated "
        "probability and derives humanized percentage. Step 5: API returns all scores and top "
        "similar sources with links. Step 6: User downloads IEEE-style DOCX report.",
    )

    add_section(
        doc,
        "4. Plagiarism Detection Algorithm",
        "The system uses sentence-transformers/all-MiniLM-L6-v2 to encode text chunks and candidate "
        "source abstracts into embeddings. Cosine similarity is computed for each chunk-source pair. "
        "Candidate source score is the average chunk similarity. Final plagiarism percentage is based "
        "on top-K averaged similarity, clamped to 0-100%. This provides semantic plagiarism detection "
        "instead of keyword-only overlap.",
    )

    add_section(
        doc,
        "5. Humanized and AI-Detected Scoring",
        "The platform uses transformer classifiers (Hello-SimpleAI/chatgpt-detector-roberta, optional "
        "ensemble with roberta-base-openai-detector) over text chunks. AI probability per chunk is "
        "averaged, optionally weighted across models. AI Detected % = P(AI) x 100 and Humanized % = "
        "(1 - P(AI)) x 100. This ensures both metrics are mathematically consistent.",
    )

    add_section(
        doc,
        "6. Open-Source Source Retrieval",
        "To provide topic-related references, the backend queries open repositories (arXiv and DOI "
        "metadata via Crossref), deduplicates sources, and semantically reranks them against user text. "
        "Top matches are returned with title, source type, URL, and similarity percentage. If live web "
        "retrieval is unavailable, a local corpus fallback is used.",
    )

    add_section(
        doc,
        "7. Performance Optimizations",
        "The project includes fast-mode optimizations: reduced chunk counts, low external API timeouts, "
        "cached query-to-source mapping, and startup model preloading. These changes significantly reduce "
        "average analysis latency and improve UX responsiveness for typical documents.",
    )

    add_section(
        doc,
        "8. Frontend Features",
        "The UI provides a clean single-page academic workflow: file picker, textarea input, submit "
        "action, three circular gauges (Plagiarism, Humanized, AI Detected), top similar source list "
        "with clickable links, and IEEE report download button. Error messages are surfaced directly "
        "from backend responses.",
    )

    add_section(
        doc,
        "9. Backend API Endpoints",
        "POST /upload: accepts file or text and returns job_id, plagiarism_percentage, "
        "humanized_percentage, ai_generated_percentage, and top_sources. GET /download-report/{job_id}: "
        "returns generated Word report for that analysis session. GET /health: service health probe.",
    )

    add_section(
        doc,
        "10. Conclusion",
        "GH Plagiarism Checker demonstrates a practical research-integrity platform that combines "
        "semantic plagiarism scoring, AI-writing detection, and source transparency. The modular "
        "architecture allows future extension with additional retrieval providers, confidence calibration, "
        "and organization-level deployment.",
    )

    doc.add_heading("References", level=1)
    refs = [
        "N. Reimers and I. Gurevych, Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks.",
        "Hugging Face model card: Hello-SimpleAI/chatgpt-detector-roberta.",
        "Hugging Face model card: roberta-base-openai-detector.",
        "Scikit-learn documentation: cosine similarity.",
        "FastAPI, React, Vite, PyPDF2, python-docx documentation.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    base = Path(__file__).resolve().parent.parent
    out_file = base / "reports" / "GH_Plagiarism_Checker_IEEE_Project_Report.docx"
    build_report(out_file)
    print(out_file)
